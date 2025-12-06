from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Title
title = doc.add_heading('Advanced Google Alerts Configuration Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('DLP/DSPM Competitive Intelligence System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(102, 126, 234)

doc.add_paragraph()

# Introduction
intro = doc.add_heading('Overview', 1)
doc.add_paragraph(
    'This guide provides a systematic approach to monitoring the Data Loss Prevention (DLP) and '
    'Data Security Posture Management (DSPM) competitive landscape using Google Alerts. By following '
    'these configurations, you\'ll receive curated intelligence on competitor activities, market trends, '
    'funding events, and product launches.'
)

# Section 1: Why Google Alerts
section1 = doc.add_heading('Why Google Alerts?', 1)
doc.add_paragraph(
    'Google Alerts is a free, powerful tool for automated competitive intelligence:'
)

benefits = doc.add_paragraph(style='List Bullet')
benefits.add_run('Real-time notifications').bold = True
benefits.add_run(' when new content matches your search criteria')

benefits = doc.add_paragraph(style='List Bullet')
benefits.add_run('Customizable frequency').bold = True
benefits.add_run(' (as-it-happens, daily, or weekly digests)')

benefits = doc.add_paragraph(style='List Bullet')
benefits.add_run('Free and easy to set up').bold = True
benefits.add_run(' - no subscription required')

benefits = doc.add_paragraph(style='List Bullet')
benefits.add_run('Email delivery').bold = True
benefits.add_run(' keeps intelligence in your inbox')

benefits = doc.add_paragraph(style='List Bullet')
benefits.add_run('Comprehensive coverage').bold = True
benefits.add_run(' across news, blogs, forums, and publications')

# Section 2: Alert Configuration Strategy
section2 = doc.add_heading('Alert Configuration Strategy', 1)
doc.add_paragraph(
    'The key to effective competitive intelligence is strategic search query construction. '
    'Below are pre-configured alerts organized by intelligence category.'
)

# Category 1: Company-Specific Monitoring
cat1 = doc.add_heading('Category 1: Company-Specific Monitoring', 2)
doc.add_paragraph('Track individual competitors for product launches, leadership changes, and company news.')

# Nightfall AI
nightfall_head = doc.add_paragraph()
nightfall_head.add_run('Nightfall AI').bold = True
nightfall_head.add_run(' (Primary Competition)')

alert1 = doc.add_paragraph(style='List Number')
alert1_run = alert1.add_run('Search Query: ')
alert1_run.bold = True
alert1.add_run('"Nightfall AI" OR "Nightfall.ai" -jobs -careers -hiring -"we\'re hiring"')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Frequency: ').bold = True
p.add_run('As-it-happens')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Sources: ').bold = True
p.add_run('News + Blogs')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Why: ').bold = True
p.add_run('Monitor product releases (like Nyx), funding, partnerships, customer wins')

# Varonis
varonis_head = doc.add_paragraph()
varonis_head.add_run('Varonis').bold = True
varonis_head.add_run(' (Legacy Leader)')

alert2 = doc.add_paragraph(style='List Number')
alert2_run = alert2.add_run('Search Query: ')
alert2_run.bold = True
alert2.add_run('Varonis ("data security" OR DLP OR DSPM) -jobs -stock -earnings -shares -"stock price"')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Frequency: ').bold = True
p.add_run('Daily')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Sources: ').bold = True
p.add_run('News only')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Why: ').bold = True
p.add_run('Track product updates, earnings (public company), strategic moves')

# Cyera
cyera_head = doc.add_paragraph()
cyera_head.add_run('Cyera').bold = True
cyera_head.add_run(' (Modern DSPM)')

alert3 = doc.add_paragraph(style='List Number')
alert3_run = alert3.add_run('Search Query: ')
alert3_run.bold = True
alert3.add_run('Cyera ("data security" OR DSPM) -jobs -careers -hiring')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Frequency: ').bold = True
p.add_run('As-it-happens')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Sources: ').bold = True
p.add_run('News + Blogs')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Why: ').bold = True
p.add_run('Recently raised $300M - monitor aggressive expansion, acquisitions, market messaging')

# Additional competitors
doc.add_paragraph()
additional = doc.add_paragraph()
additional.add_run('Additional Competitors to Monitor:').bold = True

competitors_list = [
    ('Nightfall AI', '"Nightfall AI" (DLP OR "data loss prevention") -jobs'),
    ('Varonis', 'Varonis (DLP OR "data loss prevention") -jobs'),
    ('Forcepoint', 'Forcepoint (DLP OR "data loss prevention") -jobs'),
    ('DTEX', 'DTEX (DLP OR "data loss prevention") -jobs'),
    ('Digital Guardian', '"Digital Guardian" (DLP OR "data loss prevention") -jobs'),
    ('Microsoft Purview', '"Microsoft Purview" (DLP OR "data loss prevention") -jobs'),
    ('Proofpoint', 'Proofpoint (DLP OR "data loss prevention") -jobs'),
    ('Symantec DLP', '"Symantec DLP" (DLP OR "data loss prevention") -jobs'),
    ('Cyera', 'Cyera (DSPM OR "data security") -jobs'),
    ('BigID', 'BigID (DSPM OR "data security") -jobs'),
    ('Sentra', 'Sentra (DSPM OR "data security") -jobs'),
    ('Normalyze', 'Normalyze (DSPM OR "data security") -jobs'),
    ('Laminar', 'Laminar (DSPM OR "data security") -jobs'),
    ('Dig Security', '"Dig Security" (DSPM OR "data security") -jobs'),
    ('Polar Security', '"Polar Security" (DSPM OR "data security") -jobs'),
    ('Securiti.ai', 'Securiti.ai (DSPM OR "data security") -jobs'),
    ('Netskope', 'Netskope (CASB OR SASE OR DLP) -jobs'),
    ('Zscaler', 'Zscaler (CASB OR SASE OR DLP) -jobs'),
    ('Skyhigh Security', '"Skyhigh Security" (CASB OR SASE OR DLP) -jobs'),
    ('Defensx', 'Defensx ("browser security" OR DLP) -jobs'),
    ('Ermes', 'Ermes ("browser security" OR DLP) -jobs'),
    ('SquareX', 'SquareX ("browser security" OR DLP) -jobs'),
    ('Island', 'Island ("browser security" OR DLP) -jobs'),
    ('Talon', 'Talon ("browser security" OR DLP) -jobs'),
    ('LayerX Security', '"LayerX Security" ("browser security" OR DLP) -jobs'),
    ('Perception Point', '"Perception Point" ("browser security" OR DLP) -jobs'),
    ('Redaccess', 'Redaccess ("browser security" OR DLP) -jobs'),
    ('Seraphic Security', '"Seraphic Security" ("browser security" OR DLP) -jobs'),
    ('Surf Security', '"Surf Security" ("browser security" OR DLP) -jobs'),
    ('Menlo Security', '"Menlo Security" ("browser security" OR DLP) -jobs'),
    ('Automation Anywhere', '"Automation Anywhere" (workflow OR BPM OR automation) -jobs'),
    ('UiPath', 'UiPath (workflow OR BPM OR automation) -jobs'),
    ('Blue Prism', '"Blue Prism" (workflow OR BPM OR automation) -jobs'),
    ('Microsoft Power Automate', '"Microsoft Power Automate" (workflow OR BPM OR automation) -jobs'),
    ('Camunda', 'Camunda (workflow OR BPM OR automation) -jobs'),
    ('ProcessMaker', 'ProcessMaker (workflow OR BPM OR automation) -jobs'),
    ('Appian', 'Appian (workflow OR BPM OR automation) -jobs'),
    ('Pega', 'Pega (workflow OR BPM OR automation) -jobs'),
    ('Nintex', 'Nintex (workflow OR BPM OR automation) -jobs'),
    ('Bizagi', 'Bizagi (workflow OR BPM OR automation) -jobs'),
    ('Flowable', 'Flowable (workflow OR BPM OR automation) -jobs'),
    ('Workato', 'Workato (workflow OR BPM OR automation) -jobs'),
]

for comp, query in competitors_list:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{comp}: ').bold = True
    p.add_run(query)

# Category 2: Market Intelligence
cat2 = doc.add_heading('Category 2: Market & Trend Intelligence', 2)
doc.add_paragraph('Broader market trends, emerging threats, and industry developments.')

trend_alerts = [
    ('Shadow AI Security',
     '"Shadow AI" OR "AI data leakage" OR "GenAI security" OR "LLM data loss"',
     'As-it-happens',
     'This is THE emerging trend - critical for positioning'),

    ('DLP Market Trends',
     '"data loss prevention" ("market" OR "trend" OR "forecast" OR "growth")',
     'Daily',
     'Analyst reports, market size updates, growth projections'),

    ('DSPM Evolution',
     'DSPM OR "data security posture" ("trend" OR "adoption" OR "market")',
     'Daily',
     'Track DSPM category evolution and adoption rates'),

    ('Data Breach News',
     '"data breach" (leaked OR exposed OR stolen) -bitcoin -cryptocurrency -crypto -ransomware',
     'Daily',
     'Real-world incidents drive DLP buying behavior - filters crypto noise'),

    ('AI Security & Guardrails',
     '"AI security" OR "LLM security" OR "prompt injection" OR "AI guardrails"',
     'As-it-happens',
     'Emerging AI security threats and protection mechanisms'),

    ('CASB/SASE Market',
     'CASB OR SASE ("market" OR "adoption" OR "trend") security',
     'Weekly',
     'Track broader security platform trends that include DLP'),
]

for name, query, freq, why in trend_alerts:
    alert_head = doc.add_paragraph()
    alert_head.add_run(name).bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Query: ').bold = True
    p.add_run(query)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frequency: ').bold = True
    p.add_run(freq)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Why: ').bold = True
    p.add_run(why)
    
    doc.add_paragraph()

# Category 3: Funding & M&A
cat3 = doc.add_heading('Category 3: Funding & M&A Activity', 2)
doc.add_paragraph('Track capital raises, acquisitions, and strategic investments in the space.')

funding_alerts = [
    ('Cybersecurity Funding - DLP/DSPM Focus',
     '("Series A" OR "Series B" OR "Series C" OR "Series D") (DLP OR DSPM OR "data security") cybersecurity',
     'Weekly',
     'Identify well-funded competitors and emerging players'),
    
    ('Security Acquisitions',
     'acquired (DLP OR DSPM OR "data security" OR "cloud security") cybersecurity',
     'As-it-happens',
     'M&A activity signals market consolidation and strategic shifts'),
    
    ('Venture Capital - Data Security',
     '"venture capital" OR "VC investment" ("data security" OR DSPM OR DLP)',
     'Weekly',
     'Track investor sentiment and capital flow into category'),
]

for name, query, freq, why in funding_alerts:
    alert_head = doc.add_paragraph()
    alert_head.add_run(name).bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Query: ').bold = True
    p.add_run(query)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frequency: ').bold = True
    p.add_run(freq)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Why: ').bold = True
    p.add_run(why)
    
    doc.add_paragraph()

# Category 4: Product & Technology
cat4 = doc.add_heading('Category 4: Product & Technology Developments', 2)
doc.add_paragraph('Monitor product launches, feature releases, and technological innovations.')

product_alerts = [
    ('AI-Powered Security Tools',
     '"AI-powered" OR "machine learning" (DLP OR "data security" OR "threat detection")',
     'Daily',
     'Track AI/ML integration in competitive products'),
    
    ('Cloud DLP Innovations',
     '"cloud DLP" OR "SaaS DLP" ("launch" OR "release" OR "announcement")',
     'As-it-happens',
     'New cloud-native capabilities from competitors'),
    
    ('Zero Trust + Data Security',
     '"zero trust" ("data security" OR DLP OR DSPM)',
     'Weekly',
     'Monitor convergence of zero trust and data security'),
]

for name, query, freq, why in product_alerts:
    alert_head = doc.add_paragraph()
    alert_head.add_run(name).bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Query: ').bold = True
    p.add_run(query)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frequency: ').bold = True
    p.add_run(freq)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Why: ').bold = True
    p.add_run(why)
    
    doc.add_paragraph()

# Setup Instructions
setup = doc.add_heading('Setup Instructions', 1)

doc.add_paragraph(
    'Follow these steps to configure your competitive intelligence alert system:'
)

step1 = doc.add_paragraph('Visit Google Alerts', style='List Number')
step1.add_run('\nhttps://www.google.com/alerts')

step2 = doc.add_paragraph('For Each Alert Above:', style='List Number')
step2.add_run('\n• Copy the search query exactly as written\n• Paste into "Create an alert about..." field\n• Click "Show options" to configure:\n  - How often: Set per recommendation (As-it-happens, Daily, or Weekly)\n  - Sources: News and/or Blogs (as recommended)\n  - Language: English\n  - Region: Any region (or United States for US-focused intel)\n  - How many: "Only the best results" (reduces noise)\n  - Deliver to: Your email address')

step3 = doc.add_paragraph('Create a Gmail Filter (Optional)', style='List Number')
step3.add_run('\n• Search: from:(googlealerts-noreply@google.com)\n• Create filter\n• Apply label: "Competitive Intelligence"\n• This keeps alerts organized and searchable')

step4 = doc.add_paragraph('Weekly Review Process', style='List Number')
step4.add_run('\n• Set 30-minute calendar block every Monday\n• Review all alerts from previous week\n• Update "News & Updates" tab in tracker spreadsheet\n• Identify action items (battlecard updates, talking points, etc.)')

# Advanced Tips
tips = doc.add_heading('Advanced Tips & Best Practices', 1)

tip_list = [
    ('Use Negative Keywords', 
     'Add "-jobs -careers -hiring" to exclude recruitment noise'),
    
    ('Combine Related Terms', 
     'Use OR operator: "Nightfall AI" OR "Nightfall.ai" captures variations'),
    
    ('Exclude Unrelated Content', 
     'For public companies, use "-stock -shares -earnings" to filter financial noise'),
    
    ('Use Quotes for Exact Phrases', 
     '"Shadow AI" in quotes ensures exact match, not separate words'),
    
    ('Monitor Analyst Firms', 
     'Add: Gartner OR Forrester OR "IDC" to catch analyst reports'),
    
    ('Track Executive Movements', 
     'Create alerts for: [Competitor] ("CEO" OR "CTO" OR "CMO" OR "hired" OR "appointed")'),
    
    ('Regional Intelligence', 
     'Set region to "United Kingdom" or "Australia" to monitor international expansion'),
]

for tip_title, tip_content in tip_list:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(tip_title + ': ').bold = True
    p.add_run(tip_content)

# Troubleshooting
trouble = doc.add_heading('Troubleshooting Common Issues', 1)

issues = [
    ('Too Much Noise', 
     '• Add more negative keywords (-jobs, -careers, -stock)\n• Change "How many" to "Only the best results"\n• Switch from "As-it-happens" to "Daily" digest'),
    
    ('Missing Important News', 
     '• Check if your query is too restrictive\n• Add OR variations of company/product names\n• Use broader terms (DLP instead of specific product names)'),
    
    ('Irrelevant Results', 
     '• Use quotes for exact phrases\n• Add more context terms: "Nightfall AI" cybersecurity\n• Review and refine queries monthly'),
]

for issue_title, issue_content in issues:
    p = doc.add_paragraph()
    p.add_run(issue_title).bold = True
    doc.add_paragraph(issue_content)

# Quick Start Summary
summary = doc.add_heading('Quick Start: Top 5 Must-Have Alerts', 1)
doc.add_paragraph(
    'If you\'re setting up for the first time, start with these five essential alerts:'
)

must_have = [
    '"Nightfall AI" OR "Nightfall.ai" -jobs -careers -hiring',
    '"Shadow AI" OR "AI data leakage" OR "GenAI security"',
    'Cyera ("data security" OR DSPM) -jobs -careers',
    '("Series B" OR "Series C" OR "Series D") (DLP OR DSPM OR "data security") cybersecurity',
    '"data breach" (leaked OR exposed OR stolen) -bitcoin -crypto',
]

for i, alert in enumerate(must_have, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(alert)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('—————')
footer_run.font.color.rgb = RGBColor(102, 126, 234)

doc.add_paragraph()
final = doc.add_paragraph('This alert system will provide you with continuous competitive intelligence, '
                         'enabling you to stay ahead of market trends, anticipate competitor moves, and '
                         'identify emerging threats and opportunities in the DLP/DSPM space.')
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_run = final.runs[0]
final_run.font.italic = True
final_run.font.size = Pt(10)

doc.save('google_alerts_configuration_guide.docx')
print("Google Alerts Configuration Guide created successfully!")
print("File saved: google_alerts_configuration_guide.docx")
print(f"Total competitor alerts configured: {len(competitors_list)}")
